import os
import random
import string
import json
from docx import Document
from PIL import Image
import numpy as np
import csv

# 文件夹路径
output_folder = os.path.join(os.path.dirname(__file__), "data")

# 创建文件夹（如果不存在）
os.makedirs(output_folder, exist_ok=True)

# 通用文本文件生成函数
def generate_complex_text_file(file_path, content=None):
    with open(file_path, 'w') as f:
        if content:
            f.write(content)
        else:
            random_text = ''.join(random.choices(string.ascii_letters + string.digits, k=1024))
            f.write(random_text)

# 生成复杂JSON文件
def generate_complex_json_file(file_path):
    complex_data = {"id": random.randint(1, 1000), "name": ''.join(random.choices(string.ascii_letters, k=10))}
    with open(file_path, 'w') as f:
        json.dump(complex_data, f, indent=4)

# 生成复杂CSV文件
def generate_complex_csv_file(file_path):
    with open(file_path, 'w', newline='') as csvfile:
        writer = csv.writer(csvfile)
        writer.writerow(['Name', 'Age', 'City'])
        for _ in range(100):
            writer.writerow([random.choice(['Alice', 'Bob', 'Charlie']), random.randint(20, 50), random.choice(['New York', 'Los Angeles', 'Chicago'])])

# 生成复杂Python文件
def generate_complex_python_file(file_path):
    functions = ""
    for i in range(5):  # 生成5个随机函数
        func_name = ''.join(random.choices(string.ascii_lowercase, k=8))
        param_name = ''.join(random.choices(string.ascii_lowercase, k=5))
        functions += f"def {func_name}({param_name}):\n"
        functions += f"    print('{func_name} was called with {param_name}')\n\n"
    with open(file_path, 'w') as f:
        f.write(functions)

# 生成复杂图像文件
def generate_complex_image_file(file_path, width=256, height=256):
    array = np.random.rand(height, width, 3) * 255  # 随机生成像素点
    img = Image.fromarray(array.astype('uint8')).convert('RGB')
    img.save(file_path)

# 生成复杂Word文件
def generate_complex_word_file(file_path):
    doc = Document()
    doc.add_heading('Complex Word Document', 0)
    for i in range(5):
        doc.add_paragraph('Random paragraph ' + ''.join(random.choices(string.ascii_letters, k=50)))
    doc.save(file_path)

# 生成HTML文件
def generate_complex_html_file(file_path):
    html_content = f"<html><body><h1>Sample HTML {random.randint(1, 100)}</h1><p>This is a sample HTML file.</p></body></html>"
    generate_complex_text_file(file_path, html_content)

# 生成其他文件格式
def generate_binary_file(file_path, size_kb=10):
    with open(file_path, 'wb') as f:
        f.write(os.urandom(size_kb * 1024))

# 生成每种文件类型的50个文件
def generate_50_files_per_type(output_folder):
    file_types = {
        "txt": generate_complex_text_file,
        "md": generate_complex_text_file,
        "csv": generate_complex_csv_file,
        "json": generate_complex_json_file,
        "xml": lambda path: generate_complex_text_file(path, "<root><item>Sample</item></root>"),
        "html": generate_complex_html_file,
        "log": lambda path: generate_complex_text_file(path, "Log entry: " + ''.join(random.choices(string.ascii_letters, k=100))),
        "ini": lambda path: generate_complex_text_file(path, "[settings]\noption=value"),
        "yml": lambda path: generate_complex_text_file(path, "key: value"),
        "py": generate_complex_python_file,
        "java": lambda path: generate_complex_text_file(path, "public class Sample {}"),
        "cpp": lambda path: generate_complex_text_file(path, "#include <iostream>\nint main() { return 0; }"),
        "js": lambda path: generate_complex_text_file(path, "console.log('Hello, world');"),
        "php": lambda path: generate_complex_text_file(path, "<?php echo 'Hello, PHP'; ?>"),
        "ts": lambda path: generate_complex_text_file(path, "let example: string = 'Hello, TypeScript';"),
        "rb": lambda path: generate_complex_text_file(path, "puts 'Hello, Ruby'"),
        "go": lambda path: generate_complex_text_file(path, "package main\nfunc main() { println('Hello, Go') }"),
        "kt": lambda path: generate_complex_text_file(path, "fun main() { println('Hello, Kotlin') }"),
        "dart": lambda path: generate_complex_text_file(path, "void main() { print('Hello, Dart'); }"),
        "docx": generate_complex_word_file,
        "pdf": lambda path: generate_binary_file(path, 10),
        "pptx": lambda path: generate_binary_file(path, 10),
        "xls": lambda path: generate_binary_file(path, 10),
        "xlsx": lambda path: generate_binary_file(path, 10),
        "rtf": lambda path: generate_binary_file(path, 10),
        "png": generate_complex_image_file,
        "jpg": generate_complex_image_file,
        "gif": generate_complex_image_file,
        "ico": generate_complex_image_file,
        "mp3": lambda path: generate_binary_file(path, 10),
        "wav": lambda path: generate_binary_file(path, 10),
        "mp4": lambda path: generate_binary_file(path, 10),
        "avi": lambda path: generate_binary_file(path, 10),
        "mov": lambda path: generate_binary_file(path, 10),
        "mkv": lambda path: generate_binary_file(path, 10),
        "zip": lambda path: generate_binary_file(path, 10),
        "tar": lambda path: generate_binary_file(path, 10),
        "rar": lambda path: generate_binary_file(path, 10),
        "exe": lambda path: generate_binary_file(path, 10),
        "bat": lambda path: generate_complex_text_file(path, "@echo off\necho Hello, world"),
        "sh": lambda path: generate_complex_text_file(path, "echo Hello, world"),
        "bin": lambda path: generate_binary_file(path, 10),
        "iso": lambda path: generate_binary_file(path, 10),
        "class": lambda path: generate_binary_file(path, 10),
        "h": lambda path: generate_complex_text_file(path, "#ifndef SAMPLE_H\n#define SAMPLE_H\n#endif"),
        "swift": lambda path: generate_complex_text_file(path, "print('Hello, Swift')"),
        "scss": lambda path: generate_complex_text_file(path, "$color: #333;\nbody { color: $color; }")
    }

    # 对每种类型生成50个文件
    for ext, func in file_types.items():
        for i in range(50):
            file_name = f"file_{i+1}.{ext}"
            file_path = os.path.join(output_folder, ext, file_name)
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
            func(file_path)

# 生成数据集
generate_50_files_per_type(output_folder)